import os
from pathlib import Path
from docx import Document

#method to extract text from .txt or .docx
def file_text_to_string(file_path):
	file_path = Path(file_path)

	#handle bad file path
	if not file_path.exists():
		raise FileNotFoundError(f"File {file_path} not found.")

	#normalize suffix to lower case
	ext =file_path.suffix.lower()
	#if .txt
	if ext==".txt":
		return file_path.read_text(encoding="utf-8")
	#if .docx scan all text 
	elif ext == ".docx":
		doc = Document(file_path)
		#container for text
		text=[]
		#read regular text from file
		for paragraph in doc.paragraphs:
			text.append(paragraph.text)
		#read text from embedded tables
		for table in doc.tables:
			for row in table.rows:
				for cell in row.cells:
					text.append(cell.text)
		return '\n'.join(text)
	else:
		raise ValueError(f"Unsupported file type: {file_path.suffix}")

